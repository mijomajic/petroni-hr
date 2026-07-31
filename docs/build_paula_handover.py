from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(__file__).with_name("paula-technical-handover-2026-07-31.md")
OUTPUT = Path(__file__).with_name("Petroni-tehnicki-handover-2026-07-31.docx")
LOGO = ROOT / "static" / "brand" / "petroni-logo.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "2B2B2B"
MUTED = "6B7178"
LIGHT = "F2F4F7"
GOLD = "F5C518"
WHITE = "FFFFFF"


def set_font(run, *, name="Calibri", size=11, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_border(paragraph, *, left=None, bottom=None):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side, config in (("left", left), ("bottom", bottom)):
        if not config:
            continue
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(config.get("size", 12)))
        border.set(qn("w:space"), str(config.get("space", 6)))
        border.set(qn("w:color"), config.get("color", BLUE))
        borders.append(border)


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url, *, bold=False, italic=False, size=11, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(size_node)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text, *, size=11, color=INK, bold=False):
    cursor = 0
    for match in TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Courier New", size=max(8.5, size - 1), color=DARK_BLUE, bold=bold)
        else:
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url, bold=bold, size=size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=size, color=color, bold=bold)


def add_numbering_definition(document, *, ordered):
    numbering = document.part.numbering_part.element
    existing_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Stranica ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def set_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    values = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in values.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_table(document, rows):
    columns = len(rows[0])
    widths_by_count = {
        2: [2700, 6660],
        3: [1900, 3460, 4000],
        4: [1800, 2300, 2200, 3060],
    }
    widths = widths_by_count.get(columns)
    if rows[0][0] == "Naziv / obrazac":
        widths = [2300, 2200, 2300, 2560]
    elif rows[0][0] == "Servis":
        widths = [1600, 2500, 2400, 2860]
    if widths is None:
        base = 9360 // columns
        widths = [base] * columns
        widths[-1] += 9360 - sum(widths)

    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    configure_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_inline(paragraph, value, size=9, color=INK, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, LIGHT)
        if row_index == 0:
            set_repeat_table_header(table.rows[row_index])
    configure_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def clean_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    set_styles(document)
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Quiet multi-page running furniture from the standard business brief preset.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("PETRONI  |  TEHNIČKI HANDOVER  |  31. 7. 2026.")
    set_font(run, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    document.core_properties.title = "Petroni — tehnički pregled sustava i plan za puštanje u rad"
    document.core_properties.subject = "Tehnički handover i go-live pregled"
    document.core_properties.author = "Petroni"
    document.core_properties.keywords = "Petroni, handover, hosting, Supabase, Vercel, CorvusPay"

    # Memo masthead cover.
    if LOGO.exists():
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_paragraph.paragraph_format.space_after = Pt(30)
        logo = logo_paragraph.add_run().add_picture(str(LOGO), width=Inches(2.45))
        logo._inline.docPr.set("name", "Petroni logo")
        logo._inline.docPr.set("descr", "Žuti Petroni logotip")

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("TEHNIČKI HANDOVER I GO-LIVE AUDIT")
    set_font(run, size=10, bold=True, color=DARK_BLUE)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Petroni sustav")
    set_font(run, size=27, bold=True, color=INK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("Arhitektura, vlasništvo, privatnost, oporavak i plan za sigurno puštanje u rad")
    set_font(run, size=14, color=MUTED)

    metadata = [
        ("Datum pregleda", "31. srpnja 2026."),
        ("Status", "Radni handover za vlasnika sustava i IT podršku"),
        ("Opseg", "Javna stranica, booking, webshop, administracija, podaci, plaćanje, hosting i održavanje"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{label}: ")
        set_font(run, size=10.5, bold=True, color=INK)
        run = paragraph.add_run(value)
        set_font(run, size=10.5, color=INK)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(16)
    paragraph_border(rule, bottom={"size": 18, "space": 3, "color": GOLD})

    quote = next(line[2:].strip() for line in lines if line.startswith("> "))
    callout = document.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(6)
    paragraph_shading(callout, "F8FAFC")
    paragraph_border(callout, left={"size": 18, "space": 8, "color": BLUE})
    add_inline(callout, quote, size=10.5, color=INK, bold=False)

    footer_note = document.add_paragraph()
    footer_note.paragraph_format.space_before = Pt(26)
    footer_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer_note.add_run("Radni dokument — bez tajnih vrijednosti i bez tvrdnje o pravnoj usklađenosti.")
    set_font(run, size=9, italic=True, color=MUTED)
    document.add_page_break()

    # Body begins at the first numbered section.
    start = next(index for index, line in enumerate(lines) if line.startswith("## 1."))
    index = start
    active_list_kind = None
    active_list_id = None
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            active_list_kind = None
            active_list_id = None
            index += 1
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline(paragraph, line[3:].strip(), size=16, color=BLUE, bold=True)
            active_list_kind = None
            index += 1
            continue
        if line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline(paragraph, line[4:].strip(), size=13, color=BLUE, bold=True)
            active_list_kind = None
            index += 1
            continue
        if line.startswith("#### "):
            paragraph = document.add_paragraph(style="Heading 3")
            add_inline(paragraph, line[5:].strip(), size=12, color=DARK_BLUE, bold=True)
            active_list_kind = None
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1]):
            table_lines = [line]
            index += 2  # Skip delimiter row.
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, [clean_table_row(row) for row in table_lines])
            active_list_kind = None
            continue
        if line.startswith("- "):
            if active_list_kind != "bullet":
                active_list_id = add_numbering_definition(document, ordered=False)
                active_list_kind = "bullet"
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
            apply_numbering(paragraph, active_list_id)
            add_inline(paragraph, line[2:].strip())
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            if active_list_kind != "number":
                active_list_id = add_numbering_definition(document, ordered=True)
                active_list_kind = "number"
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
            apply_numbering(paragraph, active_list_id)
            add_inline(paragraph, re.sub(r"^\d+\. ", "", line))
            index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph_shading(paragraph, "F8FAFC")
            paragraph_border(paragraph, left={"size": 16, "space": 8, "color": BLUE})
            paragraph.paragraph_format.left_indent = Inches(0.16)
            add_inline(paragraph, line[2:].strip(), size=10.5)
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline(paragraph, line)
        active_list_kind = None
        active_list_id = None
        index += 1

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
